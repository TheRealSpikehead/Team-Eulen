import os
from pathlib import Path
from math import modf
from datetime import datetime, timedelta

from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker, scoped_session

from data_access.data_base import init_db  # Sicherstellen, dass diese Implementierung verfügbar ist
from data_models.models import Room, Guest, Booking, RegisteredGuest  # Sicherstellen, dass diese Klassen verfügbar sind
from docx import Document
import tkinter as tk
from tkinter import messagebox, filedialog


class ReservationManager:
    def __init__(self, db_file):
        if not db_file.is_file():
            init_db(str(db_file), generate_example_data=True)
        self._engine = create_engine(f'sqlite:///{db_file}')
        self._session = scoped_session(sessionmaker(bind=self._engine))

    def make_booking(self, room_id: int, guest_id: int, number_of_guests: int, start_date: datetime, end_date: datetime,
                     comment: str = None):
        query_room = select(Room).where(Room.id == room_id)
        room = self._session.execute(query_room).scalars().one()
        query_guest = select(Guest).where(Guest.id == guest_id)
        guest = self._session.execute(query_guest).scalars().one()
        if number_of_guests <= room.max_guests:
            new_booking = Booking(
                room=room,
                guest=guest,
                number_of_guests=number_of_guests,
                start_date=start_date,
                end_date=end_date,
                comment=comment
            )
            self._session.add(new_booking)
            self._session.commit()
            return new_booking
        else:
            raise ValueError('Number of guests is bigger than the room allows')

    def get_price(self, booking: Booking):
        duration = booking.end_date - booking.start_date
        price = booking.room.price * duration.days
        price_mwst = round(price + price * 0.077, 2)
        return self.round_for_currency(price), self.round_for_currency(price_mwst)

    def round_for_currency(self, price):
        first_digit = modf(price)[0] * 10
        to_round = round(modf(first_digit)[0] / 10, 2)
        if to_round < 0.03:
            diff = round(0.0 - to_round, 2)
        elif to_round < 0.07:
            diff = round(0.05 - to_round, 2)
        else:
            diff = 0.1 - to_round
        result = round(price + diff, 2)
        return result

    def get_bookings(self, registered_guest: RegisteredGuest):
        query = select(Booking).where(Booking.guest == registered_guest)
        result = self._session.execute(query).scalars().all()
        return result

    def get_filtered_rooms(self, number_of_guests):
        query = select(Room).where(Room.max_guests >= number_of_guests)
        return self._session.execute(query).scalars().all()


def update_room_menu():
    try:
        number_of_guests = int(entry_number_of_guests.get())
        filtered_rooms = manager.get_filtered_rooms(number_of_guests)
        room_choices = {room.id: f"Raum {room.id} (Max Gäste: {room.max_guests})" for room in filtered_rooms}
        room_var.set(next(iter(room_choices)))  # Setze die Standardauswahl auf den ersten Raum

        # Aktualisiere das Dropdown-Menü
        room_menu['menu'].delete(0, 'end')
        for room_id in room_choices.keys():
            room_menu['menu'].add_command(label=room_choices[room_id], command=tk._setit(room_var, room_id))
    except ValueError:
        messagebox.showerror("Fehler", "Bitte eine gültige Anzahl von Gästen eingeben")


def create_booking():
    try:
        start_date = datetime.strptime(entry_start_date.get(), '%d.%m.%Y')
        duration = int(entry_duration.get())
        end_date = start_date + timedelta(days=duration)
        number_of_guests = int(entry_number_of_guests.get())
        comment = entry_comment.get()
        room_id = int(room_var.get())
        guest_id = int(guest_var.get())

        booking = manager.make_booking(room_id, guest_id, number_of_guests, start_date, end_date, comment)

        # Überprüfen, ob ein Word-Dokument erstellt werden soll
        if create_doc_var.get() == 1:
            # Word-Dokument erstellen und Speicherort auswählen
            file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
            if file_path:
                doc = Document()
                doc.add_heading('Buchungsbestätigung', 0)

                doc.add_paragraph(f'Startdatum: {start_date.strftime("%d.%m.%Y")}')
                doc.add_paragraph(f'Enddatum: {end_date.strftime("%d.%m.%Y")}')
                doc.add_paragraph(f'Anzahl der Gäste: {number_of_guests}')
                doc.add_paragraph(f'Kommentar: {comment}')
                doc.add_paragraph(f'Raum ID: {room_id}')
                doc.add_paragraph(f'Gast ID: {guest_id}')

                doc.save(file_path)
                messagebox.showinfo("Erfolg", f"Buchung erfolgreich erstellt! Dokument gespeichert als {file_path}")
        else:
            messagebox.showinfo("Erfolg", "Buchung erfolgreich erstellt!")

    except Exception as e:
        messagebox.showerror("Fehler", str(e))


if __name__ == '__main__':
    db_path = Path("../data/database.db")
    manager = ReservationManager(db_path)

    root = tk.Tk()
    root.title("Reservation Manager")

    # Eingabefelder und Labels
    tk.Label(root, text="Startdatum (DD.MM.YYYY):").pack(pady=5)
    entry_start_date = tk.Entry(root)
    entry_start_date.pack(pady=5)

    tk.Label(root, text="Dauer (Tage):").pack(pady=5)
    entry_duration = tk.Entry(root)
    entry_duration.pack(pady=5)

    tk.Label(root, text="Anzahl der Gäste:").pack(pady=5)
    entry_number_of_guests = tk.Entry(root)
    entry_number_of_guests.pack(pady=5)

    tk.Label(root, text="Kommentar:").pack(pady=5)
    entry_comment = tk.Entry(root)
    entry_comment.pack(pady=5)

    # Dropdown-Menü für Räume
    tk.Label(root, text="Raum:").pack(pady=5)
    room_var = tk.StringVar(root)
    room_menu = tk.OptionMenu(root, room_var, "")
    room_menu.pack(pady=5)

    # Aktualisiere die Raum-Dropdown-Liste, wenn die Anzahl der Gäste eingegeben wird
    entry_number_of_guests.bind("<FocusOut>", lambda event: update_room_menu())

    # Dropdown-Menü für Gäste
    tk.Label(root, text="Gast:").pack(pady=5)
    guest_var = tk.StringVar(root)
    guests = manager._session.execute(select(Guest)).scalars().all()
    guest_choices = {guest.id: guest for guest in guests}
    guest_var.set(next(iter(guest_choices)))  # Setze die Standardauswahl auf den ersten Gast
    guest_menu = tk.OptionMenu(root, guest_var, *guest_choices.keys())
    guest_menu.pack(pady=5)

    # Checkbox für Word-Dokument-Erstellung
    create_doc_var = tk.IntVar()
    create_doc_check = tk.Checkbutton(root, text="Word-Dokument erstellen", variable=create_doc_var)
    create_doc_check.pack(pady=5)

    # Buchungs-Button
    button = tk.Button(root, text="Buchung erstellen", command=create_booking)
    button.pack(pady=20)

    root.mainloop()
